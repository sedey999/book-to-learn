#!/usr/bin/env python3
"""
Multi-format text extraction for book-to-learn.
Extracts plain text from PDF/DOCX/HTML/EPUB/TXT/MD/RTF with fallback chains.

Usage:
  python extract_text.py <file_path> [--out <output.txt>]

Output: writes extracted text to <output.txt> (default: stdout).
Single-source failures raise ExtractionError but print partial results.
"""
import os, sys, re, subprocess, html, argparse

class ExtractionError(Exception):
    pass

# ── PDF ──
def extract_pdf(path):
    """PDF: pdftotext → pypdf → pdfminer.six"""
    # 1. pdftotext (poppler) — fastest, best for text-heavy PDFs
    try:
        r = subprocess.run(['pdftotext', '-layout', path, '-'],
                           capture_output=True, text=True, timeout=120)
        if r.returncode == 0 and len(r.stdout.strip()) > 100:
            return r.stdout
    except (FileNotFoundError, subprocess.TimeoutExpired):
        pass
    # 2. pypdf
    try:
        from pypdf import PdfReader
        reader = PdfReader(path)
        parts = []
        for page in reader.pages:
            t = page.extract_text() or ''
            parts.append(t)
        text = '\n\n'.join(parts)
        if len(text.strip()) > 100:
            return text
    except Exception:
        pass
    # 3. pdfminer.six
    try:
        from pdfminer.high_level import extract_text
        text = extract_text(path)
        if text and len(text.strip()) > 100:
            return text
    except Exception:
        pass
    raise ExtractionError(f"PDF extraction failed (all methods): {path}")

# ── DOCX ──
def extract_docx(path):
    """DOCX: python-docx → stdlib zipfile/xml"""
    try:
        import docx
        doc = docx.Document(path)
        parts = [p.text for p in doc.paragraphs if p.text.strip()]
        # also extract tables
        for table in doc.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells]
                parts.append('\t'.join(cells))
        text = '\n'.join(parts)
        if text.strip():
            return text
    except Exception:
        pass
    # fallback: stdlib zipfile
    try:
        import zipfile, xml.etree.ElementTree as ET
        with zipfile.ZipFile(path) as z:
            with z.open('word/document.xml') as f:
                tree = ET.parse(f)
        ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
        texts = []
        for t in tree.iter('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t'):
            if t.text:
                texts.append(t.text)
        text = ''.join(texts)
        if text.strip():
            return text
    except Exception:
        pass
    raise ExtractionError(f"DOCX extraction failed: {path}")

# ── HTML ──
def extract_html(path):
    """HTML: beautifulsoup4 → stdlib html.parser"""
    try:
        from bs4 import BeautifulSoup
        with open(path, 'r', encoding='utf-8', errors='ignore') as f:
            soup = BeautifulSoup(f.read(), 'html.parser')
        # remove scripts/styles
        for tag in soup(['script', 'style', 'nav', 'footer', 'header']):
            tag.decompose()
        text = soup.get_text(separator='\n', strip=True)
        if text.strip():
            return text
    except Exception:
        pass
    # fallback: stdlib
    try:
        from html.parser import HTMLParser
        class TextExtractor(HTMLParser):
            def __init__(self):
                super().__init__(); self.text = []; self.skip = False
            def handle_starttag(self, tag, attrs):
                if tag in ('script', 'style'): self.skip = True
            def handle_endtag(self, tag):
                if tag in ('script', 'style'): self.skip = False
                if tag in ('p', 'div', 'br', 'li', 'h1','h2','h3','h4','h5','h6'):
                    self.text.append('\n')
            def handle_data(self, data):
                if not self.skip: self.text.append(data)
        with open(path, 'r', encoding='utf-8', errors='ignore') as f:
            p = TextExtractor(); p.feed(f.read())
        text = re.sub(r'\n{3,}', '\n\n', ''.join(p.text)).strip()
        if text:
            return text
    except Exception:
        pass
    raise ExtractionError(f"HTML extraction failed: {path}")

# ── EPUB ──
def extract_epub(path):
    """EPUB: ebooklib + bs4 → stdlib zipfile"""
    try:
        import ebooklib
        from ebooklib import epub
        from bs4 import BeautifulSoup
        book = epub.read_epub(path, options={'ignore_ncx': True})
        parts = []
        for item in book.get_items_of_type(ebooklib.ITEM_DOCUMENT):
            soup = BeautifulSoup(item.get_content(), 'html.parser')
            for tag in soup(['script', 'style']):
                tag.decompose()
            text = soup.get_text(separator='\n', strip=True)
            if text.strip():
                parts.append(text)
        text = '\n\n'.join(parts)
        if text.strip():
            return text
    except Exception:
        pass
    # fallback: stdlib zipfile
    try:
        import zipfile, html as html_mod
        parts = []
        with zipfile.ZipFile(path) as z:
            for name in z.namelist():
                if name.endswith(('.html', '.xhtml', '.htm')):
                    with z.open(name) as f:
                        raw = f.read().decode('utf-8', errors='ignore')
                        text = re.sub(r'<[^>]+>', ' ', raw)
                        text = html_mod.unescape(text)
                        text = re.sub(r'\s+', ' ', text).strip()
                        if text:
                            parts.append(text)
        text = '\n\n'.join(parts)
        if text.strip():
            return text
    except Exception:
        pass
    raise ExtractionError(f"EPUB extraction failed: {path}")

# ── TXT/MD/RST/AsciiDoc ──
def extract_plain(path):
    """TXT/MD/RST/ADOC: direct read"""
    for enc in ['utf-8', 'gbk', 'gb2312', 'big5', 'latin-1']:
        try:
            with open(path, 'r', encoding=enc) as f:
                text = f.read()
            if text.strip():
                return text
        except (UnicodeDecodeError, LookupError):
            continue
    raise ExtractionError(f"Plain text extraction failed: {path}")

# ── RTF ──
def extract_rtf(path):
    """RTF: striprtf → regex fallback"""
    try:
        from striprtf.striprtf import rtf_to_text
        with open(path, 'r', encoding='utf-8', errors='ignore') as f:
            text = rtf_to_text(f.read())
        if text.strip():
            return text
    except ImportError:
        pass
    # regex fallback: strip RTF control words
    try:
        with open(path, 'r', encoding='utf-8', errors='ignore') as f:
            raw = f.read()
        text = re.sub(r'\\[a-z]+-?\d* ?', '', raw)
        text = re.sub(r'[{}]', '', text)
        text = re.sub(r'\\\*\\', '', text)
        text = html.unescape(text)
        text = re.sub(r'\s+', ' ', text).strip()
        if text:
            return text
    except Exception:
        pass
    raise ExtractionError(f"RTF extraction failed: {path}")

# ── dispatcher ──
EXTRACTORS = {
    '.pdf': extract_pdf,
    '.docx': extract_docx,
    '.html': extract_html, '.htm': extract_html,
    '.epub': extract_epub,
    '.txt': extract_plain, '.md': extract_plain, '.markdown': extract_plain,
    '.rst': extract_plain, '.adoc': extract_plain, '.asciidoc': extract_plain,
    '.rtf': extract_rtf,
}

def extract(path):
    """Extract text from a file. Returns text string. Raises ExtractionError on failure."""
    if not os.path.exists(path):
        raise ExtractionError(f"File not found: {path}")
    ext = os.path.splitext(path)[1].lower()
    fn = EXTRACTORS.get(ext)
    if not fn:
        raise ExtractionError(f"Unsupported format: {ext} ({path})")
    return fn(path)

def main():
    ap = argparse.ArgumentParser(description='Multi-format text extraction')
    ap.add_argument('file', help='file path to extract')
    ap.add_argument('--out', help='output file path (default: stdout)')
    args = ap.parse_args()
    try:
        text = extract(args.file)
        if args.out:
            with open(args.out, 'w', encoding='utf-8') as f:
                f.write(text)
            print(json.dumps({'ok': True, 'out': args.out, 'chars': len(text)},
                             ensure_ascii=False))
        else:
            sys.stdout.write(text)
    except ExtractionError as e:
        print(json.dumps({'ok': False, 'error': str(e)}, ensure_ascii=False), file=sys.stderr)
        sys.exit(1)

if __name__ == '__main__':
    import json
    main()
